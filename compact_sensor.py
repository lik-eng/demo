"""
传感器期末押题版 → 终极冲刺版 精简脚本
读取"传感器期末押题版 - 副本.docx"，按规则过滤后生成精简版。
"""
import os
import re
import docx
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# ========== 配置 ==========
SRC_DIR = r"c:\Users\lik_6\OneDrive\桌面\demo"
OUTPUT_PATH = r"C:\Users\lik_6\OneDrive\桌面\传感器期末终极冲刺版.docx"

# 核心原理关键词（这些题目即使没有★也保留）
CORE_KEYWORDS = [
    "应变效应", "惠斯通电桥", "变极距", "变面积", "电涡流",
    "压电效应", "电荷放大器", "霍尔效应", "热电偶", "冷端补偿",
    "Pt100", "外光电效应", "内光电效应", "光生伏特效应",
]

# 简答题保留编号（原文件中的题号）
SA_KEEP_NUMS = {1, 2, 3, 4, 5, 6, 7, 8, 9, 11, 12, 13, 15, 16, 18, 19, 22}


# ========== 辅助函数 ==========

def find_source_file():
    """查找源文件（传感器期末押题版 - 副本.docx）"""
    for f in os.listdir(SRC_DIR):
        if "押题" in f and "副本" in f and f.endswith(".docx"):
            return os.path.join(SRC_DIR, f)
    raise FileNotFoundError("找不到源文件「传感器期末押题版 - 副本.docx」")


def is_empty_para(para):
    """判断段落是否为空（无文本或仅空白）"""
    return not para.text.strip()


def is_heading(para):
    """判断段落是否为标题样式"""
    return para.style and "Heading" in (para.style.name or "")


def get_heading_text(para):
    """获取标题段落的纯净文本"""
    return para.text.strip()


def para_has_star(para):
    """检查段落是否以★开头"""
    return para.text.strip().startswith("★")


def para_has_core_keyword(para):
    """检查段落是否包含核心原理关键词"""
    text = para.text
    for kw in CORE_KEYWORDS:
        if kw in text:
            return True
    return False


def should_keep_question(question_lines):
    """
    判断一组题目段落是否应该保留。
    question_lines: 题目的所有段落（题干 + 选项 + 答案行等）
    返回 True 表示保留。
    """
    if not question_lines:
        return False
    # 检查题干行（第一个非空段落）
    first_text = question_lines[0].text.strip()
    # 条件1：题干以★开头
    if first_text.startswith("★"):
        return True
    # 条件2：包含核心原理关键词（检查所有行）
    for para in question_lines:
        if para_has_core_keyword(para):
            return True
    return False


def extract_sa_number(para):
    """
    从简答题题干中提取题号。
    例如 "★ 1. xxx" → 1, "15. xxx" → 15
    """
    text = para.text.strip()
    # 去掉开头的★和空格
    text = re.sub(r"^★\s*", "", text)
    # 匹配 "数字." 或 "数字、" 或 "数字．"
    m = re.match(r"(\d+)\s*[\.\．、]", text)
    if m:
        return int(m.group(1))
    return None


def copy_paragraph_formatting(src_para, dest_para):
    """复制段落级别格式"""
    # 对齐方式
    dest_para.alignment = src_para.alignment
    # 段落间距
    pf = dest_para.paragraph_format
    pf_src = src_para.paragraph_format
    if pf_src.space_before is not None:
        pf.space_before = pf_src.space_before
    if pf_src.space_after is not None:
        pf.space_after = pf_src.space_after
    if pf_src.first_line_indent is not None:
        pf.first_line_indent = pf_src.first_line_indent
    if pf_src.left_indent is not None:
        pf.left_indent = pf_src.left_indent


def copy_run_formatting(src_run, dest_run):
    """复制文本运行级别格式"""
    if src_run.bold is not None:
        dest_run.bold = src_run.bold
    if src_run.italic is not None:
        dest_run.italic = src_run.italic
    if src_run.underline is not None:
        dest_run.underline = src_run.underline
    if src_run.font.size is not None:
        dest_run.font.size = src_run.font.size
    if src_run.font.name is not None:
        dest_run.font.name = src_run.font.name
    # 复制颜色
    try:
        if src_run.font.color and src_run.font.color.rgb is not None:
            dest_run.font.color.rgb = src_run.font.color.rgb
    except Exception:
        pass


def copy_paragraph(src_para, dest_doc):
    """
    将源段落完整复制到目标文档，保留文本格式。
    返回新创建的段落对象。
    """
    new_para = dest_doc.add_paragraph()
    # 复制段落格式
    copy_paragraph_formatting(src_para, new_para)
    # 复制每个 run（保留字体、大小、颜色、粗体等）
    for run in src_para.runs:
        new_run = new_para.add_run(run.text)
        copy_run_formatting(run, new_run)
    return new_para


def copy_paragraph_with_new_text(src_para, dest_doc, new_text):
    """
    复制段落格式，但替换文本内容。
    用于重新编号简答题等场景。
    """
    new_para = dest_doc.add_paragraph()
    copy_paragraph_formatting(src_para, new_para)
    # 保持第一个 run 的格式，替换文本
    if src_para.runs:
        new_run = new_para.add_run(new_text)
        copy_run_formatting(src_para.runs[0], new_run)
    else:
        new_para.add_run(new_text)
    return new_para


def add_title_page(doc):
    """添加文档标题页"""
    # 空行留白
    doc.add_paragraph()
    doc.add_paragraph()

    # 主标题：黑体三号（16pt），居中
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(12)
    run = title.add_run("《传感器原理及检测技术》期末终极冲刺版")
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    run.font.size = Pt(16)
    run.bold = True

    # 副标题
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(6)
    run2 = subtitle.add_run("（精简版）")
    run2.font.name = "黑体"
    run2._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    run2.font.size = Pt(14)

    # 说明文字：宋体小四（12pt）
    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    desc.paragraph_format.space_after = Pt(24)
    run3 = desc.add_run("适用专科期末，已剔除冗余与低频考点，只留必考核心")
    run3.font.name = "宋体"
    run3._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run3.font.size = Pt(12)
    run3.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # 分隔线
    sep = doc.add_paragraph()
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sep.paragraph_format.space_after = Pt(18)
    run4 = sep.add_run("—" * 40)
    run4.font.size = Pt(10)
    run4.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)


def add_section_heading(doc, text, level=1):
    """
    添加章节标题。
    level=1 → 黑体三号（类似 Heading1）
    level=2 → 黑体小三号（类似 Heading2）
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(text)
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    if level == 1:
        run.font.size = Pt(16)
        run.bold = True
    else:
        run.font.size = Pt(14)
        run.bold = True
    return p


def group_paragraphs_by_empty_line(paragraphs):
    """
    将段落按空行分组。
    每组是一段连续的非空段落。
    返回 list of list of paragraph objects。
    """
    groups = []
    current_group = []

    for para in paragraphs:
        if is_empty_para(para):
            if current_group:
                groups.append(current_group)
                current_group = []
        else:
            current_group.append(para)

    # 最后一批
    if current_group:
        groups.append(current_group)

    return groups


def renumber_question(para, new_num):
    """
    给题干段落重新编号。
    保留★标记，替换原题号为新的序号。
    """
    text = para.text.strip()
    # 匹配模式：可选★ + 数字 + 标点 + 剩余内容
    # 替换为题号
    has_star = text.startswith("★")
    # 去掉原编号
    if has_star:
        text = text[1:].strip()  # 去掉★
    # 去掉原题号（匹配开头的数字+标点）
    text = re.sub(r"^\d+\s*[\.\．、]\s*", "", text)
    # 重新组装
    if has_star:
        return f"★ {new_num}. {text}"
    else:
        return f"{new_num}. {text}"


# ========== 主处理逻辑 ==========

def main():
    print("[1/5] 正在打开源文件...")
    src_path = find_source_file()
    print(f"  源文件: {src_path}")
    src_doc = Document(src_path)
    total = len(src_doc.paragraphs)
    print(f"  总段落数: {total}")

    # ===== 第一步：扫描文档结构，动态定位各区域边界 =====
    print("[2/5] 扫描文档结构...")

    # 先找出所有标题段落（仅 Heading 样式的段落）
    all_headings = []  # [(index, style_name, text)]
    for i, p in enumerate(src_doc.paragraphs):
        if is_heading(p):
            all_headings.append((i, p.style.name, get_heading_text(p)))

    print(f"  共找到 {len(all_headings)} 个标题段落:")
    for idx, style, text in all_headings:
        # 安全打印（避免emoji编码错误）
        print(f"    [{idx:5d}] [{style:15s}] {text[:60]}")

    # 基于标题文本定位各区域
    # 约束：只匹配 Heading 1 中的主要区域标题
    boundaries = {}
    mock_indices = []

    # 判断是否为一级/二级标题（兼容 "Heading 1" / "Heading1" / "标题 1" 等变体）
    def is_h1(style_name):
        return "heading 1" in style_name.lower() or "标题 1" in style_name

    def is_h2(style_name):
        return "heading 2" in style_name.lower() or "标题 2" in style_name

    for idx, style, text in all_headings:
        # 主要区域标题
        if is_h1(style):
            if "200" in text and "选" in text:
                boundaries["mc_start"] = idx
            elif "100" in text and "判断" in text:
                boundaries["tf_start"] = idx
            elif "100" in text and "填空" in text:
                boundaries["fb_start"] = idx
            elif "30" in text and "简答" in text:
                boundaries["sa_start"] = idx
            elif "模拟" in text and ("五套" in text or "试卷" in text):
                boundaries["mock_heading"] = idx

        # 模拟卷子标题
        if is_h2(style):
            if "模拟卷" in text:
                mock_indices.append(idx)

    # 如果仍未检测到模拟卷，降级到宽泛匹配
    if not mock_indices:
        for idx, style, text in all_headings:
            if "模拟卷" in text and "heading" in style.lower():
                mock_indices.append(idx)

    print(f"\n  区域边界: {boundaries}")
    print(f"  模拟卷索引: {mock_indices} (共{len(mock_indices)}套)")

    # 构建保留区域列表
    keep_ranges = []

    # MC 区域
    mc_h = boundaries.get("mc_start")
    tf_h = boundaries.get("tf_start")
    fb_h = boundaries.get("fb_start")
    sa_h = boundaries.get("sa_start")
    mock_h = boundaries.get("mock_heading")

    # 保留 MC 标题 + 内容
    if mc_h is not None:
        keep_ranges.append((mc_h, mc_h + 1, "mc_heading"))
        # MC 内容：从标题下一段到 TF 标题前（或到文档预估位置）
        mc_end = tf_h if tf_h is not None else 1753
        keep_ranges.append((mc_h + 1, mc_end, "mc_body"))

    # 保留 TF 标题 + 内容
    if tf_h is not None:
        keep_ranges.append((tf_h, tf_h + 1, "tf_heading"))
        tf_end = fb_h if fb_h is not None else 1955
        keep_ranges.append((tf_h + 1, tf_end, "tf_body"))

    # 保留 FB 标题 + 内容
    if fb_h is not None:
        keep_ranges.append((fb_h, fb_h + 1, "fb_heading"))
        fb_end = sa_h if sa_h is not None else 2158
        keep_ranges.append((fb_h + 1, fb_end, "fb_body"))

    # 保留 SA 标题 + 内容
    if sa_h is not None:
        keep_ranges.append((sa_h, sa_h + 1, "sa_heading"))
        sa_end = mock_h if mock_h is not None else 2251
        keep_ranges.append((sa_h + 1, sa_end, "sa_body"))

    # 保留模拟卷标题
    if mock_h is not None:
        keep_ranges.append((mock_h, mock_h + 1, "mock_heading"))

    # 保留模拟卷1、3、5（按索引取第0、2、4个，对应原卷1/3/5）
    if len(mock_indices) >= 5:
        # 模拟卷1: 索引0 到 索引1
        keep_ranges.append((mock_indices[0], mock_indices[1], "mock_1"))
        # 模拟卷3: 索引2 到 索引3
        keep_ranges.append((mock_indices[2], mock_indices[3], "mock_3"))
        # 模拟卷5: 索引4 到文档末尾
        keep_ranges.append((mock_indices[4], total, "mock_5"))
    elif len(mock_indices) >= 3:
        # 备用：如果只检测到3个，全部保留
        print("  [警告] 仅检测到部分模拟卷，尝试全部保留")
        for i, mi in enumerate(mock_indices):
            end_idx = mock_indices[i + 1] if i + 1 < len(mock_indices) else total
            keep_ranges.append((mi, end_idx, f"mock_{i+1}"))

    # 按起始索引排序
    keep_ranges.sort(key=lambda x: x[0])

    print(f"\n  保留区域数: {len(keep_ranges)}")
    for start, end, label in keep_ranges:
        print(f"    [{start:5d}-{end:5d}] {label} ({end-start}段)")

    # ===== 第二步：逐区域处理 =====
    print("[3/5] 开始过滤处理...")

    # 创建新文档
    new_doc = Document()

    # 设置默认字体
    style = new_doc.styles["Normal"]
    font = style.font
    font.name = "宋体"
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    # 设置页面边距（A4常规）
    for section in new_doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    # 添加标题页
    add_title_page(new_doc)

    # 统计变量
    mc_count = 0
    tf_count = 0
    fb_count = 0
    sa_count = 0

    # 逐区域处理
    for start, end, label in keep_ranges:
        region_paras = list(src_doc.paragraphs[start:end])

        if label == "mc_heading":
            add_section_heading(new_doc, "一、选择题（高频核心精选）", level=1)

        elif label == "mc_body":
            # 将MC段落按空行分组，每组一道题
            groups = group_paragraphs_by_empty_line(region_paras)
            kept_groups = []
            for g in groups:
                if should_keep_question(g):
                    kept_groups.append(g)
            mc_count = len(kept_groups)
            print(f"  选择题: 保留 {mc_count}/{len(groups)} 题")

            # 重新编号并复制
            for idx, group in enumerate(kept_groups, 1):
                # 第一行是题干，重新编号
                new_text = renumber_question(group[0], idx)
                copy_paragraph_with_new_text(group[0], new_doc, new_text)
                # 剩余行直接复制
                for para in group[1:]:
                    copy_paragraph(para, new_doc)
                # 添加空行分隔
                new_doc.add_paragraph()

        elif label == "tf_heading":
            add_section_heading(new_doc, "二、判断题（高频核心精选）", level=1)

        elif label == "tf_body":
            groups = group_paragraphs_by_empty_line(region_paras)
            kept_groups = []
            for g in groups:
                if should_keep_question(g):
                    kept_groups.append(g)
            tf_count = len(kept_groups)
            print(f"  判断题: 保留 {tf_count}/{len(groups)} 题")

            for idx, group in enumerate(kept_groups, 1):
                new_text = renumber_question(group[0], idx)
                copy_paragraph_with_new_text(group[0], new_doc, new_text)
                for para in group[1:]:
                    copy_paragraph(para, new_doc)
                new_doc.add_paragraph()

        elif label == "fb_heading":
            add_section_heading(new_doc, "三、填空题（高频核心精选）", level=1)

        elif label == "fb_body":
            groups = group_paragraphs_by_empty_line(region_paras)
            kept_groups = []
            for g in groups:
                if should_keep_question(g):
                    kept_groups.append(g)
            fb_count = len(kept_groups)
            print(f"  填空题: 保留 {fb_count}/{len(groups)} 题")

            for idx, group in enumerate(kept_groups, 1):
                new_text = renumber_question(group[0], idx)
                copy_paragraph_with_new_text(group[0], new_doc, new_text)
                for para in group[1:]:
                    copy_paragraph(para, new_doc)
                new_doc.add_paragraph()

        elif label == "sa_heading":
            add_section_heading(new_doc, "四、简答题（必背核心）", level=1)

        elif label == "sa_body":
            groups = group_paragraphs_by_empty_line(region_paras)
            kept_groups = []
            for g in groups:
                if not g:
                    continue
                # 在组内查找题号（可能在第一个非描述段落中）
                num = None
                q_idx = 0  # 题干在组内的起始位置
                for j, p in enumerate(g):
                    num = extract_sa_number(p)
                    if num is not None:
                        q_idx = j
                        break
                # 跳过无法识别题号的组（如纯说明行）
                if num is None:
                    continue
                if num in SA_KEEP_NUMS:
                    # 只保留题干及其之后的内容，跳过前面的描述
                    kept_groups.append((num, g[q_idx:]))
            # 按原题号排序
            kept_groups.sort(key=lambda x: x[0])
            sa_count = len(kept_groups)
            print(f"  简答题: 保留 {sa_count}/30 题 (编号: {[n for n, _ in kept_groups]})")

            for new_idx, (orig_num, group) in enumerate(kept_groups, 1):
                # 重新编号题干
                new_question_text = renumber_question(group[0], new_idx)
                copy_paragraph_with_new_text(group[0], new_doc, new_question_text)
                # 复制答案行
                for para in group[1:]:
                    copy_paragraph(para, new_doc)
                new_doc.add_paragraph()

        elif label == "mock_heading":
            add_section_heading(new_doc, "五、模拟试卷（精选3套）", level=1)

        elif label in ("mock_1", "mock_3", "mock_5"):
            # 模拟卷直接复制（保留原始格式和内部分节结构）
            mock_num = label[-1]
            print(f"  复制模拟卷{mock_num}: {len(region_paras)} 段")
            add_section_heading(new_doc, f"模拟卷{mock_num}", level=2)
            for para in region_paras:
                if is_heading(para):
                    # 跳过原始标题（已用 add_section_heading 替代）
                    continue
                elif not is_empty_para(para):
                    copy_paragraph(para, new_doc)
                else:
                    new_doc.add_paragraph()

    # ===== 第三步：保存 =====
    print("[4/5] 正在保存文档...")
    new_doc.save(OUTPUT_PATH)
    print(f"  已保存到: {OUTPUT_PATH}")

    # ===== 第四步：验证 =====
    print("[5/5] 验证结果:")
    print(f"  - 选择题: {mc_count} 题")
    print(f"  - 判断题: {tf_count} 题")
    print(f"  - 填空题: {fb_count} 题")
    print(f"  - 简答题: {sa_count} 题")
    print(f"  - 模拟卷: 3 套（原卷1/3/5）")
    print(f"\n  总计约 {mc_count + tf_count + fb_count + sa_count} 题 + 3套模拟卷")
    print(f"  原版 ~800题 -> 精简版 ~{mc_count + tf_count + fb_count + sa_count + 3*45}题")

    # 检查文件大小
    size_kb = os.path.getsize(OUTPUT_PATH) / 1024
    print(f"  文件大小: {size_kb:.0f} KB")
    print("\n[DONE] 请用 Word 打开查看。")


if __name__ == "__main__":
    main()
