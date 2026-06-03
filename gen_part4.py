"""
gen_part4.py — 传感器期末押题版 第4步
合并Part1-Part3的docx文件，添加答案汇总表，生成最终文档

运行：python gen_part4.py
输出：传感器期末押题版.docx
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
import copy
import os

BASE = r"c:\Users\lik_6\OneDrive\桌面\demo"
OUTPUT = os.path.join(BASE, "传感器期末押题版.docx")

PART1 = os.path.join(BASE, "传感器期末押题_part1.docx")
PART2 = os.path.join(BASE, "传感器期末押题_part2.docx")
PART3 = os.path.join(BASE, "传感器期末押题_part3.docx")


def add_page_break(doc):
    """添加分页符"""
    from docx.oxml import OxmlElement
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._element.append(br)
    return p


def merge_docx(source_path, dest_doc):
    """将source_path中的内容追加到dest_doc末尾（保留格式）"""
    if not os.path.exists(source_path):
        print(f"   ⚠ 文件不存在: {source_path}")
        return

    src = Document(source_path)

    # 复制每个段落
    for element in src.element.body:
        # 跳过第一节的sectPr
        if element.tag.endswith('}sectPr'):
            continue
        dest_doc.element.body.append(copy.deepcopy(element))


def add_answer_summary(doc):
    """添加答案汇总表"""
    add_page_break(doc)

    # 标题
    h = doc.add_heading('标准答案速查表', level=1)
    for run in h.runs:
        run.font.name = '黑体'

    doc.add_paragraph('说明：以下汇总各题型答案，方便快速核对。★为高频考点。')

    # ===== 第一章专项答案 =====
    doc.add_heading('第一章 绪论 专项练习答案', level=2)

    # 选择题答案
    doc.add_paragraph('一、选择题（1-10题）', style='List Bullet')
    table = doc.add_table(rows=3, cols=10, style='Table Grid')
    answers_mc = ['B', 'D', 'C', 'A', 'B', 'C', 'A', 'B', 'D', 'A']
    for i, a in enumerate(answers_mc):
        table.cell(0, i).text = str(i + 1)
        table.cell(1, i).text = a
    # 合并第二行为答案行
    for cell in table.rows[1].cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
                run.font.bold = True
    doc.add_paragraph()

    # 判断题答案
    doc.add_paragraph('二、判断题（11-20题）：√ × √ × √ × √ √ × √', style='List Bullet')

    # 填空题答案
    doc.add_paragraph('三、填空题（21-30题）：敏感元件 | 生物量 | 智能 | 智能化 | 输出信号 | 有源 | 应用 | 敏感 | 物性 | 通信', style='List Bullet')

    doc.add_paragraph()

    # ===== 第2-5章答案 =====
    doc.add_heading('第二章 传感器基本特性答案', level=2)
    doc.add_paragraph('选择：1.A 2.B 3.C 4.B 5.A', style='List Bullet')
    doc.add_paragraph('判断：6.√ 7.√ 8.√ 9.× 10.√', style='List Bullet')
    doc.add_paragraph('填空：11.线性度 12.输出量 13.时间 14.阻尼 15.线性', style='List Bullet')

    doc.add_heading('第三章 电阻式传感器答案', level=2)
    doc.add_paragraph('选择：1.B 2.C 3.A 4.D 5.B', style='List Bullet')
    doc.add_paragraph('判断：6.√ 7.× 8.√ 9.× 10.×', style='List Bullet')
    doc.add_paragraph('填空：11.电阻率 12.线性 13.平整光洁 14.位移 15.低', style='List Bullet')

    doc.add_heading('第四章 电容式传感器答案', level=2)
    doc.add_paragraph('选择：1.A 2.D 3.C 4.B 5.D', style='List Bullet')
    doc.add_paragraph('判断：6.√ 7.× 8.× 9.√ 10.√', style='List Bullet')
    doc.add_paragraph('填空：11.介电 12.变介质 13.交流电桥 14.线性 15.高压', style='List Bullet')

    doc.add_heading('第五章 电感式传感器答案', level=2)
    doc.add_paragraph('选择：1.D 2.A 3.B 4.C 5.A', style='List Bullet')
    doc.add_paragraph('判断：6.√ 7.× 8.× 9.× 10.√', style='List Bullet')
    doc.add_paragraph('填空：11.电涡流 12.线圈 13.位移 14.互感 15.工作', style='List Bullet')

    # ===== 第6-9章答案 =====
    doc.add_heading('第六章 压电式传感器答案', level=2)
    doc.add_paragraph('选择：1.A 2.D 3.B 4.A 5.C', style='List Bullet')
    doc.add_paragraph('判断：6.√ 7.× 8.√ 9.√ 10.×', style='List Bullet')
    doc.add_paragraph('填空：11.电介质 12.电压 13.动 14.石英 15.并', style='List Bullet')

    doc.add_heading('第七章 霍尔传感器答案', level=2)
    doc.add_paragraph('选择：1.A 2.B 3.D 4.A 5.B', style='List Bullet')
    doc.add_paragraph('判断：6.√ 7.× 8.√ 9.× 10.√', style='List Bullet')
    doc.add_paragraph('填空：11.霍尔灵敏 12.III-V 13.非接触 14.磁场 15.磁感应', style='List Bullet')

    doc.add_heading('第八章 热电式传感器答案', level=2)
    doc.add_paragraph('选择：1.A 2.C 3.B 4.A 5.B', style='List Bullet')
    doc.add_paragraph('判断：6.× 7.√ 8.× 9.√ 10.×', style='List Bullet')
    doc.add_paragraph('填空：11.热电（塞贝克） 12.温差 13.100 14.负 15.冰点', style='List Bullet')

    doc.add_heading('第九章 光电式传感器答案', level=2)
    doc.add_paragraph('选择：1.C 2.B 3.A 4.B 5.C', style='List Bullet')
    doc.add_paragraph('判断：6.√ 7.× 8.√ 9.× 10.√', style='List Bullet')
    doc.add_paragraph('填空：11.外光电 12.光电导 13.反向 14.光生伏特 15.二次电子', style='List Bullet')

    add_page_break(doc)

    # ===== 200道选择题答案汇总 =====
    doc.add_heading('200道选择题答案速查', level=2)

    # 按章节排列
    mc_answers = {
        "第1章(1-25)": "1.B 2.C 3.A 4.D 5.C 6.D 7.B 8.A 9.B 10.D 11.C 12.B 13.D 14.D 15.A 16.C 17.A 18.B 19.C 20.A 21.B 22.D 23.A 24.D 25.B",
        "第2章(26-47)": "26.C 27.B 28.A 29.D 30.B 31.C 32.A 33.C 34.B 35.A 36.B 37.D 38.A 39.D 40.B 41.A 42.B 43.C 44.A 45.B 46.D 47.C",
        "第3章(48-69)": "48.B 49.C 50.A 51.D 52.B 53.C 54.A 55.B 56.D 57.A 58.C 59.B 60.A 61.B 62.A 63.C 64.B 65.A 66.C 67.B 68.D 69.D",
        "第4章(70-91)": "70.A 71.B 72.C 73.A 74.B 75.C 76.A 77.D 78.A 79.D 80.C 81.B 82.C 83.D 84.B 85.A 86.A 87.B 88.C 89.A 90.D 91.B",
        "第5章(92-113)": "92.A 93.B 94.A 95.D 96.B 97.C 98.B 99.D 100.A 101.D 102.B 103.A 104.C 105.A 106.B 107.A 108.C 109.D 110.D 111.B 112.C 113.A",
        "第6章(114-135)": "114.C 115.B 116.C 117.A 118.B 119.A 120.A 121.D 122.B 123.A 124.B 125.A 126.B 127.C 128.A 129.C 130.B 131.A 132.B 133.A 134.C 135.C",
        "第7章(136-157)": "136.B 137.D 138.A 139.B 140.C 141.B 142.A 143.C 144.B 145.A 146.C 147.B 148.A 149.A 150.A 151.D 152.B 153.C 154.A 155.B 156.D 157.B",
        "第8章(158-179)": "158.C 159.B 160.A 161.A 162.B 163.A 164.D 165.B 166.C 167.A 168.A 169.B 170.A 171.C 172.B 173.A 174.D 175.C 176.A 177.B 178.C 179.A",
        "第9章(180-200)": "180.A 181.A 182.B 183.C 184.C 185.A 186.B 187.A 188.B 189.C 190.D 191.A 192.B 193.A 194.B 195.A 196.B 197.D 198.D 199.D 200.B",
    }

    for chapter, answers in mc_answers.items():
        doc.add_paragraph(f"{chapter}：{answers}", style='List Bullet').paragraph_format.space_after = Pt(4)

    add_page_break(doc)

    # ===== 100道判断题答案 =====
    doc.add_heading('100道判断题答案速查', level=2)
    tf_answers = [
        "第1章(1-12)：1.√ 2.√ 3.× 4.√ 5.√ 6.× 7.√ 8.× 9.× 10.× 11.√ 12.√",
        "第2章(13-23)：13.√ 14.√ 15.√ 16.× 17.√ 18.× 19.× 20.√ 21.√ 22.× 23.√",
        "第3章(24-34)：24.√ 25.√ 26.× 27.√ 28.× 29.× 30.× 31.× 32.√ 33.√ 34.√",
        "第4章(35-45)：35.√ 36.× 37.√ 38.√ 39.√ 40.× 41.√ 42.√ 43.× 44.× 45.√",
        "第5章(46-56)：46.√ 47.× 48.√ 49.× 50.√ 51.√ 52.× 53.× 54.√ 55.√ 56.√",
        "第6章(57-67)：57.√ 58.× 59.√ 60.× 61.√ 62.√ 63.√ 64.× 65.√ 66.√ 67.×",
        "第7章(68-78)：68.√ 69.× 70.√ 71.√ 72.× 73.√ 74.√ 75.× 76.√ 77.√ 78.√",
        "第8章(79-89)：79.√ 80.× 81.√ 82.× 83.√ 84.× 85.√ 86.√ 87.× 88.√ 89.√",
        "第9章(90-100)：90.√ 91.× 92.× 93.√ 94.√ 95.× 96.√ 97.× 98.√ 99.√ 100.√",
    ]
    for ans in tf_answers:
        doc.add_paragraph(ans, style='List Bullet').paragraph_format.space_after = Pt(2)

    add_page_break(doc)

    # ===== 100道填空题答案 =====
    doc.add_heading('100道填空题答案速查', level=2)
    fb_chapters = [
        "第1章(1-14)：1.转换元件 2.输出信号 3.有源 4.生物量 5.模拟 6.敏感 7.电 8.应用 9.智能化 10.物性 11.处理 12.传感器 13.结构 14.湿度",
        "第2章(15-25)：15.输入 16.线性度 17.正 18.重复性 19.时间 20.阻尼 21.漂移 22.弹性 23.分辨率 24.标定 25.正弦",
        "第3章(26-36)：26.应变 27.ρL/A 28.惠斯通 29.4 30.压阻 31.横向 32.交流 33.电阻 34.线性 35.应变 36.1.7",
        "第4章(37-47)：37.εA/d 38.变介质 39.面积 40.调频 41.驱动 42.高 43.阻抗 44.差动 45.介质 46.高压 47.分布",
        "第5章(48-58)：48.电涡流 49.气隙 50.非 51.互感 52.导 53.线性可变差动变压器 54.灵敏度 55.无损 56.位移 57.减小 58.涡流",
        "第6章(59-69)：59.压电 60.电压 61.动 62.石英 63.并 64.差 65.电缆 66.后 67.超声波 68.X（电轴） 69.绝缘",
        "第7章(70-79)：70.霍尔 71.KH·I·B 72.III-V 73.非 74.位置 75.电势 76.正比 77.电流 78.优（好） 79.薄",
        "第8章(80-89)：80.热电（塞贝克） 81.温差 82.100 83.负 84.增大 85.0 86.补偿 87.有源 88.四 89.不变",
        "第9章(90-100)：90.外 91.光电导 92.反向 93.光生伏特 94.二次 95.大 96.短路 97.耦合 98.绝对 99.光栅 100.非",
    ]
    for ans in fb_chapters:
        doc.add_paragraph(ans, style='List Bullet').paragraph_format.space_after = Pt(2)

    add_page_break(doc)

    # ===== 30道简答题答案要点 =====
    doc.add_heading('30道简答题答案要点', level=2)
    sa_list = [
        "1. 传感器是感受被测量并按规律转换成可用输出信号的器件，由敏感元件+转换元件+测量电路组成。",
        "2. 按被测量分（物理/化学/生物）、按工作原理分（电阻/电容/电感等）、按能量分（有源/无源）、按输出信号分（模拟/数字）。",
        "3. 静态特性是对不随时间变化输入的响应特性，指标：灵敏度、线性度、迟滞、重复性、分辨率、精度、漂移。",
        "4. 灵敏度K=Δy/Δx，即输出变化量与输入变化量之比，反映敏感程度。",
        "5. 线性度是实际曲线与拟合直线的最大偏差，原因：原理非线性、结构非线性、弹性元件非线性。",
        "6. 迟滞是正反行程曲线不重合度，原因：弹性滞后、磁滞、摩擦和间隙。",
        "7. 应变效应：金属丝受力变形→几何尺寸变化→电阻变化，R=ρL/A，拉伸时L↑A↓R↑。",
        "8. 惠斯通电桥将微小ΔR→ΔV输出，差动半桥灵敏度×2、全桥×4，兼有温度补偿和线性改善。",
        "9. 变极距型（改d，非线性）、变面积型（改A，线性）、变介质型（改ε，测液位/厚度）。",
        "10. 优点：分辨率高、响应快、耐恶劣环境、非接触。缺点：高阻抗、易受分布电容干扰。",
        "11. 变磁阻式（改气隙→L变）、电涡流式（涡流效应→Z变）、差动变压器式（改铁芯→互感变）。",
        "12. 电涡流原理：线圈产生交变磁场→导体产生涡流→涡流反磁场使Z变。应用：位移/振动/厚度/无损检测。",
        "13. 压电效应：电介质受力变形→内部极化→表面产生电荷，Q∝F。可逆（逆压电效应=电致伸缩）。",
        "14. 要求高输入阻抗（≥10⁹Ω），防止电荷泄漏。电荷放大器优于电压放大器（电缆不敏感）。",
        "15. 静态力作用下电荷通过回路泄漏→输出衰减。适合动态测量（振动、冲击、加速度）。",
        "16. 霍尔效应：半导体在磁场中通电流→垂直方向产生UH=KHIB，与I·B乘积成正比。",
        "17. 特点：非接触、结构简单、响应快。应用：测B/I/位移/角度/转速/汽车点火/无刷电机。",
        "18. 热电偶（塞贝克效应）：两种不同导体构成回路，T≠T₀时产生热电势，查分度表得温度。",
        "19. 分度表冷端0℃，实际冷端≠0℃需补偿。方法：冰点槽、补偿导线、补偿电桥、计算修正。",
        "20. 热电阻（金属/PTC/线性好/量程宽）vs 热敏电阻（半导体/NTC或PTC/灵敏度高/非线性/量程窄）。",
        "21. 均质导体定律、中间导体定律、中间温度定律。",
        "22. 外光电效应（光电管/倍增管）、光电导效应（光敏电阻）、光生伏特效应（光电池/光电二极管）。",
        "23. 光敏电阻：半导体吸光→载流子↑→R↓。特性：光照越强R越小，暗/亮电阻比越大灵敏度越高。",
        "24. 差动优点：灵敏度加倍、消偶次非线性、共模干扰抑制、温度补偿（对称抵消）。",
        "25. 选型原则：根据对象/环境、灵敏度适中、精度够用、频响覆盖、稳定性好、考虑成本/体积/安装。",
        "26. 标定：实验确定输入-输出关系。目的：获得静态参数、确定精度/误差、检验设计指标。",
        "27. 称重传感器：弹性元件受力→应变片变形→ΔR→电桥→ΔV→放大/A/D→处理显示。",
        "28. 光电转速：码盘转动→光断续→光电元件输出脉冲→计数→n=60N/(Z·t)。",
        "29. 电压放大器：简单廉价但电缆影响大；电荷放大器：电缆不敏感、低频好，但电路更复杂。",
        "30. 趋势：微型化（MEMS）、智能化、多功能化、网络化、集成化、新材料应用。",
    ]
    for s in sa_list:
        doc.add_paragraph(s, style='List Bullet').paragraph_format.space_after = Pt(4)

    add_page_break(doc)

    # ===== 高频考点清单 =====
    doc.add_heading('★★★★★ 高频考点清单（15个，务必掌握）', level=2)
    key_points = [
        "1. 传感器组成：敏感元件 + 转换元件 + 测量电路",
        "2. 传感器分类：按被测量/工作原理/能量关系/输出信号分",
        "3. 灵敏度 K=Δy/Δx",
        "4. 线性度：实际曲线与拟合直线的最大偏差",
        "5. 迟滞：正反行程曲线不重合度（弹性滞后+摩擦）",
        "6. 重复性：同向多次测量不一致度（随机误差）",
        "7. 动态特性：阶跃响应（时间常数τ/固有频率ωn+阻尼比ξ）",
        "8. 应变效应：金属丝受力→尺寸变化→ΔR（R=ρL/A）",
        "9. 惠斯通电桥：ΔR→ΔV，单臂/半桥(×2)/全桥(×4)",
        "10. 电容式传感器原理：变极距型/变面积型/变介质型（C=εA/d）",
        "11. 电感式传感器分类：变磁阻式/电涡流式/差动变压器式",
        "12. 压电效应：电介质受力→极化→电荷，适合动态测量",
        "13. 霍尔效应：UH=KH·I·B，半导体+磁场→垂直方向电动势",
        "14. 热电偶测温原理：塞贝克效应，两种导体+T差→热电势",
        "15. 光电效应：外光电/内光电（光电导+光生伏特）",
    ]
    for kp in key_points:
        p = doc.add_paragraph(kp, style='List Bullet')
        p.paragraph_format.space_after = Pt(6)
        for run in p.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)


def main():
    print("=" * 60)
    print("传感器期末押题版 — 第4步：合并文档")
    print("=" * 60)

    # 创建最终文档
    final_doc = Document()

    # 设置默认字体
    style = final_doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # A4页面
    for section in final_doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # 合并Part1
    print("\n[1/4] 合并 Part1（封面+第1～5章专项）...")
    merge_docx(PART1, final_doc)

    # 合并Part2
    print("[2/4] 合并 Part2（第6～9章专项+200选择+100判断）...")
    merge_docx(PART2, final_doc)

    # 合并Part3
    print("[3/4] 合并 Part3（100填空+30简答+5套模拟卷）...")
    merge_docx(PART3, final_doc)

    # 添加答案汇总
    print("[4/4] 添加答案汇总表和考点清单...")
    add_answer_summary(final_doc)

    # 保存
    final_doc.save(OUTPUT)
    size_kb = os.path.getsize(OUTPUT) / 1024
    print(f"\n[完成] 第4步完成！最终文档已保存至：{OUTPUT}")
    print(f"   文件大小：{size_kb:.0f} KB")
    print(f"   包含：封面 + 全部题目 + 答案汇总 + 高频考点清单")


if __name__ == "__main__":
    main()
